"""python-docx 核心引擎：新建文档 + 模板填充"""

import os
import re
from lxml import etree

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from kai_docx_generator.styles.standard import (
    apply_standard_style,
    set_run_cjk_font,
    STANDARD_CONFIG,
)


class DocxEngine:
    """统一的 python-docx 引擎"""

    FONT = "Microsoft YaHei"

    def __init__(self, style_config: dict | None = None):
        self.style_config = style_config or STANDARD_CONFIG

    def generate_from_spec(self, spec: dict) -> Document:
        """
        新建模式：从 DocSpec AST 生成 .docx。

        Args:
            spec: DocSpec AST（来自 parser）

        Returns:
            Document 对象
        """
        doc = Document()
        apply_standard_style(doc, self.style_config)

        meta = spec.get("metadata", {})

        # 页眉
        if meta.get("header"):
            self._build_header(doc, meta["header"])

        # 页脚
        if meta.get("footer"):
            self._build_footer(doc, meta["footer"])

        # TOC（如果 metadata 中请求了）
        if meta.get("toc"):
            self._insert_toc(doc, meta["toc"])

        # 渲染 sections
        for section in spec.get("sections", []):
            self._render_section(doc, section)

        # 如果没有 sections，渲染顶层 blocks
        if not spec.get("sections"):
            for block in spec.get("blocks", []):
                self._render_block(doc, block)

        return doc

    def _render_section(self, doc: Document, section: dict) -> list:
        """渲染一个 section 及其子 section"""
        elements = []

        # 标题
        heading = section.get("heading")
        if heading and heading.get("text"):
            level = min(heading.get("level", 1), 3)
            p = doc.add_heading(heading["text"], level=level)
            for run in p.runs:
                set_run_cjk_font(run, self.FONT)
            elements.append(p)

        # blocks
        for block in section.get("blocks", []):
            elements.append(self._render_block(doc, block))

        # children
        for child in section.get("children", []):
            elements.extend(self._render_section(doc, child))

        return elements

    def _render_block(self, doc: Document, block: dict):
        """渲染单个 block"""
        block_type = block.get("type")

        if block_type == "paragraph":
            if "runs" in block:
                p = doc.add_paragraph()
                for run_data in block["runs"]:
                    run = p.add_run(run_data.get("text", ""))
                    set_run_cjk_font(run, self.FONT)
                    if run_data.get("bold"):
                        run.bold = True
                    if run_data.get("italic"):
                        run.italic = True
                return p
            else:
                p = doc.add_paragraph(block.get("text", ""))
                for run in p.runs:
                    set_run_cjk_font(run, self.FONT)
                if block.get("bold"):
                    for run in p.runs:
                        run.bold = True
                if block.get("italic"):
                    for run in p.runs:
                        run.italic = True
                return p

        elif block_type == "bullet_list":
            last_p = None
            for item in block.get("items", []):
                p = doc.add_paragraph(item, style="List Bullet")
                for run in p.runs:
                    set_run_cjk_font(run, self.FONT)
                last_p = p
            return last_p

        elif block_type == "numbered_list":
            last_p = None
            for item in block.get("items", []):
                p = doc.add_paragraph(item, style="List Number")
                for run in p.runs:
                    set_run_cjk_font(run, self.FONT)
                last_p = p
            return last_p

        elif block_type == "table":
            return self._render_table(doc, block)

        elif block_type == "callout":
            return self._render_callout(doc, block)

        elif block_type == "image":
            return self._render_image(doc, block)

        elif block_type == "page_break":
            doc.add_page_break()
            return None

        elif block_type == "hr":
            return self._render_hr(doc)

        elif block_type == "code_block":
            return self._render_code_block(doc, block)

        return None

    def _render_table(self, doc: Document, block: dict):
        """渲染表格"""
        headers = block.get("headers", [])
        rows = block.get("rows", [])
        col_count = len(headers)

        if col_count == 0:
            return None

        table = doc.add_table(rows=1 + len(rows), cols=col_count, style="Table Grid")
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # 设置表格宽度为 100% 页面宽度
        from docx.shared import Emu, Twips
        tbl_pr = table._element.tblPr
        tbl_w = etree.SubElement(tbl_pr, qn("w:tblW"))
        tbl_w.set(qn("w:w"), "5000")
        tbl_w.set(qn("w:type"), "pct")

        # 表头
        for i, header in enumerate(headers):
            cell = table.rows[0].cells[i]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(header)
            run.bold = True
            run.font.color.rgb = RGBColor.from_string("FFFFFF")
            set_run_cjk_font(run, self.FONT)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            self._set_cell_shading(cell, "2563EB")

        # 数据行（斑马纹）
        for r, row in enumerate(rows):
            for c, val in enumerate(row):
                if c >= col_count:
                    break
                cell = table.rows[r + 1].cells[c]
                cell.text = ""
                p = cell.paragraphs[0]
                run = p.add_run(self._space_emoji(str(val)))
                set_run_cjk_font(run, self.FONT)
                if r % 2 == 1:
                    self._set_cell_shading(cell, "F8FAFC")

        return table

    def _space_emoji(self, text: str) -> str:
        """在连续 emoji 之间插入空格，防止渲染重叠"""
        import re
        return re.sub(r"(\u2B50+)", lambda m: " ".join(m.group(0)), text)

    def _set_cell_shading(self, cell, color_hex: str):
        """设置单元格背景色"""
        shading = etree.SubElement(
            cell._element.get_or_add_tcPr(),
            qn("w:shd"),
        )
        shading.set(qn("w:fill"), color_hex)
        shading.set(qn("w:val"), "clear")

    def _set_cell_width(self, cell, width):
        """设置单元格宽度（单位：docx.shared 对象）"""
        from docx.shared import Twips
        tc_pr = cell._element.get_or_add_tcPr()
        tc_w = etree.SubElement(tc_pr, qn("w:tcW"))
        tc_w.set(qn("w:w"), str(int(Twips(width).twips)))
        tc_w.set(qn("w:type"), "dxa")

    def _render_callout(self, doc: Document, block: dict):
        """用单格表格模拟 callout/引用框"""
        # 先插入空段落作为间距
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_before = Pt(6)
        spacer.paragraph_format.space_after = Pt(6)

        color_map = {
            "info": {"fill": "EBF5FF", "border": "2563EB"},
            "warning": {"fill": "FFF7ED", "border": "EA580C"},
            "success": {"fill": "F0FDF4", "border": "16A34A"},
        }
        colors = color_map.get(block.get("style", "info"), color_map["info"])

        table = doc.add_table(rows=1, cols=1, style="Table Grid")
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        cell = table.rows[0].cells[0]
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(block.get("text", ""))
        set_run_cjk_font(run, self.FONT)

        # 左侧粗边框 + 背景色
        tc_pr = cell._element.get_or_add_tcPr()
        tc_borders = etree.SubElement(tc_pr, qn("w:tcBorders"))
        for side, val in [
            ("left", "12"), ("top", "4"), ("bottom", "4"), ("right", "4")
        ]:
            border = etree.SubElement(tc_borders, qn(f"w:{side}"))
            border.set(qn("w:val"), "single")
            border.set(qn("w:sz"), val)
            border.set(qn("w:color"), colors["border"])
            border.set(qn("w:space"), "0")

        self._set_cell_shading(cell, colors["fill"])
        return table

    def _render_image(self, doc: Document, block: dict):
        """渲染图片"""
        path = block.get("path", "")
        caption = block.get("caption", "")
        width = block.get("width", "100%")

        # Security: Reject path traversal attempts
        if ".." in path or path.startswith("/"):
            p = doc.add_paragraph(f"[图片加载失败: 无效路径 {path}]")
            for run in p.runs:
                set_run_cjk_font(run, self.FONT)
            return p

        if not path or not os.path.exists(path):
            p = doc.add_paragraph(f"[图片加载失败: {path}]")
            for run in p.runs:
                set_run_cjk_font(run, self.FONT)
            return p

        img_width = None
        if width and width != "100%":
            match = re.match(r"^([\d.]+)(in|cm)?$", str(width).strip())
            if match:
                try:
                    val = float(match.group(1))
                    unit = match.group(2)
                    if unit == "cm":
                        from docx.shared import Cm
                        img_width = Cm(val)
                    else:
                        img_width = Inches(val)
                except (ValueError, TypeError):
                    img_width = None

        p = doc.add_picture(path, width=img_width)
        if caption:
            caption_p = doc.add_paragraph(caption)
            caption_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in caption_p.runs:
                set_run_cjk_font(run, self.FONT)
                run.font.size = Pt(9)
                run.font.color.rgb = RGBColor.from_string("999999")
        return p

    def _render_hr(self, doc: Document):
        """渲染分割线（底部边框段落）"""
        p = doc.add_paragraph()
        pPr = p._element.get_or_add_pPr()
        pBdr = etree.SubElement(pPr, qn("w:pBdr"))
        bottom = etree.SubElement(pBdr, qn("w:bottom"))
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "12")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "CCCCCC")
        return p

    def _render_code_block(self, doc: Document, block: dict):
        """渲染代码块（等宽字体，带边框的背景框）"""
        code = block.get("code", "")
        language = block.get("language", "")

        # 先插入空段落作为间距
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_before = Pt(6)
        spacer.paragraph_format.space_after = Pt(6)

        # 用单格表格模拟代码块
        table = doc.add_table(rows=1, cols=1, style="Table Grid")
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = table.rows[0].cells[0]
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(6)

        # 语言标签
        if language:
            run = p.add_run(f"[{language}] ")
            set_run_cjk_font(run, self.FONT)
            run.font.size = Pt(8)
            run.font.color.rgb = RGBColor.from_string("999999")
            run.bold = True
            run = p.add_run("\n" + code)
        else:
            run = p.add_run(code)

        run.font.name = "Consolas"
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string("333333")

        # 设置单元格浅灰背景
        self._set_cell_shading(cell, "F5F5F5")
        return table

    def _build_header(self, doc: Document, header_meta: dict):
        """构建页眉"""
        section = doc.sections[0]
        header = section.header

        header.paragraphs[0].clear()

        p = header.paragraphs[0]
        run = p.add_run(header_meta.get("text", ""))
        set_run_cjk_font(run, self.FONT)
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string("999999")

        if header_meta.get("alignment") == "left":
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        elif header_meta.get("alignment") == "right":
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        else:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        pPr = p._element.get_or_add_pPr()
        pBdr = etree.SubElement(pPr, qn("w:pBdr"))
        bottom = etree.SubElement(pBdr, qn("w:bottom"))
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "8")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), "2563EB")

    def _insert_toc(self, doc: Document, toc_config: dict | None = None):
        """插入 TOC 目录（field code，需 Word/WPS 打开时更新域）"""
        max_level = 3
        if isinstance(toc_config, dict):
            max_level = toc_config.get("max_level", 3)

        # 分页符
        doc.add_page_break()

        # 目录标题
        p = doc.add_heading("目录", level=1)
        for run in p.runs:
            set_run_cjk_font(run, self.FONT)

        # TOC field code
        p = doc.add_paragraph()
        r = p._element.makeelement(qn("w:r"), {})
        fld_char_begin = etree.SubElement(r, qn("w:fldChar"))
        fld_char_begin.set(qn("w:fldCharType"), "begin")
        p._element.append(r)

        r2 = p._element.makeelement(qn("w:r"), {})
        instr = etree.SubElement(r2, qn("w:instrText"))
        instr.set(qn("xml:space"), "preserve")
        instr.text = f' TOC \\o "1-{max_level}" \\h \\z \\u '
        p._element.append(r2)

        r3 = p._element.makeelement(qn("w:r"), {})
        fld_char_end = etree.SubElement(r3, qn("w:fldChar"))
        fld_char_end.set(qn("w:fldCharType"), "end")
        p._element.append(r3)

        # 再分页
        doc.add_page_break()

    def _build_footer(self, doc: Document, footer_meta: dict):
        """构建页脚（含页码 field code）"""
        section = doc.sections[0]
        footer = section.footer

        footer.paragraphs[0].clear()
        p = footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

        if footer_meta.get("page_number"):
            r = p._element.makeelement(qn("w:r"), {})
            fld_char_begin = etree.SubElement(r, qn("w:fldChar"))
            fld_char_begin.set(qn("w:fldCharType"), "begin")
            p._element.append(r)

            r2 = p._element.makeelement(qn("w:r"), {})
            instr = etree.SubElement(r2, qn("w:instrText"))
            instr.set(qn("xml:space"), "preserve")
            instr.text = " PAGE "
            p._element.append(r2)

            r3 = p._element.makeelement(qn("w:r"), {})
            fld_char_end = etree.SubElement(r3, qn("w:fldChar"))
            fld_char_end.set(qn("w:fldCharType"), "end")
            p._element.append(r3)
