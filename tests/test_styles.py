"""Test that all 7 style presets can be applied without errors"""
from docx import Document
from kai_docx_generator.styles.standard import apply_standard_style
from kai_docx_generator.styles.contract import apply_contract_style
from kai_docx_generator.styles.report import apply_report_style
from kai_docx_generator.styles.meeting import apply_meeting_style
from kai_docx_generator.styles.business_letter import apply_business_letter_style
from kai_docx_generator.styles.tech_spec import apply_tech_spec_style
from kai_docx_generator.styles.dispatch import apply_dispatch_style
from kai_docx_generator.styles.notice import apply_notice_style


class TestAllStyles:

    def test_standard_style(self):
        doc = Document()
        apply_standard_style(doc)
        assert doc.styles["Normal"].font.name == "Microsoft YaHei"

    def test_contract_style(self):
        doc = Document()
        apply_contract_style(doc)
        assert doc.styles["Normal"].font.name == "FangSong"

    def test_report_style(self):
        doc = Document()
        apply_report_style(doc)
        assert doc.styles["Normal"].font.name == "Microsoft YaHei"

    def test_meeting_style(self):
        doc = Document()
        apply_meeting_style(doc)
        assert doc.styles["Normal"].font.name == "Microsoft YaHei"

    def test_business_letter_style(self):
        doc = Document()
        apply_business_letter_style(doc)
        assert doc.styles["Normal"].font.name == "Microsoft YaHei"

    def test_tech_spec_style(self):
        doc = Document()
        apply_tech_spec_style(doc)
        assert doc.styles["Normal"].font.name == "Consolas"

    def test_dispatch_style(self):
        doc = Document()
        apply_dispatch_style(doc)
        # 红头文件，H1 是红色
        from kai_docx_generator.styles.dispatch import DISPATCH_CONFIG
        assert DISPATCH_CONFIG["Heading 1"]["color"] == "FF0000"
        assert doc.styles["Normal"].font.name == "FangSong"

    def test_notice_style(self):
        doc = Document()
        apply_notice_style(doc)
        assert doc.styles["Normal"].font.name == "FangSong"
