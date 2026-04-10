"""DocxEngine tests (includes Style tests)"""

from docx import Document
from docx.document import Document as DocumentClass
from docx.shared import Pt
from docx.shared import RGBColor
from docx.oxml.ns import qn

from kai_docx_generator.styles.standard import (
    apply_standard_style,
    set_run_cjk_font,
    STANDARD_CONFIG,
)


def test_standard_style_applied():
    doc = Document()
    apply_standard_style(doc)

    normal = doc.styles["Normal"]
    assert normal.font.name == "Microsoft YaHei"
    assert normal.font.size.pt == 11


def test_heading_styles():
    doc = Document()
    apply_standard_style(doc)

    for name in ("Heading 1", "Heading 2", "Heading 3"):
        style = doc.styles[name]
        assert style.font.name == "Microsoft YaHei"


def test_heading_colors():
    doc = Document()
    apply_standard_style(doc)

    h1 = doc.styles["Heading 1"]
    assert h1.font.color.rgb == RGBColor.from_string("1A1A2E")

    h2 = doc.styles["Heading 2"]
    assert h2.font.color.rgb == RGBColor.from_string("2563EB")


def test_custom_style_config():
    """自定义配置应覆盖默认值"""
    doc = Document()
    custom = {
        "Normal": {
            "font": "Arial",
            "eastAsia": "Arial",
            "size": Pt(12),
            "color": None,
            "bold": False,
            "italic": False,
            "space_after": Pt(12),
            "line_spacing": 2.0,
        },
    }
    apply_standard_style(doc, custom)
    assert doc.styles["Normal"].font.name == "Arial"
    assert doc.styles["Normal"].font.size.pt == 12


def test_set_run_cjk_font():
    doc = Document()
    p = doc.add_paragraph("测试文本")
    run = p.runs[0]
    set_run_cjk_font(run, "SimSun")
    assert run.font.name == "SimSun"
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.get_or_add_rFonts()
    assert rFonts.get(qn("w:eastAsia")) == "SimSun"


import os
import tempfile
from lxml import etree
from docx.oxml.ns import qn
from kai_docx_generator.engine.docx_engine import DocxEngine
from kai_docx_generator.parser.md_to_ast import parse_markdown_to_docspec


class TestDocxEngineGenerate:

    def _make_doc_from_md(self, md: str):
        spec = parse_markdown_to_docspec(md)
        engine = DocxEngine()
        return engine.generate_from_spec(spec)

    def test_generate_returns_document(self):
        doc = self._make_doc_from_md("# 标题\n\n正文")
        assert isinstance(doc, DocumentClass)

    def test_heading_rendering(self):
        doc = self._make_doc_from_md("# 一级标题\n\n## 二级标题\n\n### 三级标题")
        headings = [p for p in doc.paragraphs if p.style.name.startswith("Heading")]
        assert len(headings) == 3
        assert headings[0].text == "一级标题"
        assert headings[1].text == "二级标题"
        assert headings[2].text == "三级标题"

    def test_paragraph_rendering(self):
        doc = self._make_doc_from_md("这是第一段。\n\n这是第二段。")
        paragraphs = [p for p in doc.paragraphs if p.style.name == "Normal"]
        assert len(paragraphs) == 2
        assert paragraphs[0].text == "这是第一段。"
        assert paragraphs[1].text == "这是第二段。"

    def test_bullet_list_rendering(self):
        doc = self._make_doc_from_md("- 项目一\n- 项目二\n- 项目三")
        list_paras = [p for p in doc.paragraphs if "List" in p.style.name]
        assert len(list_paras) == 3
        assert list_paras[0].text == "项目一"

    def test_numbered_list_rendering(self):
        doc = self._make_doc_from_md("1. 第一步\n2. 第二步")
        list_paras = [p for p in doc.paragraphs if "List" in p.style.name]
        assert len(list_paras) == 2

    def test_table_rendering(self):
        doc = self._make_doc_from_md("| 列1 | 列2 |\n| --- | --- |\n| A | B |")
        assert len(doc.tables) == 1
        assert doc.tables[0].rows[0].cells[0].text == "列1"
        assert doc.tables[0].rows[1].cells[0].text == "A"

    def test_callout_rendering(self):
        doc = self._make_doc_from_md("> 这是一条提示信息")
        # Callout is rendered as a table
        assert len(doc.tables) == 1
        assert "提示信息" in doc.tables[0].rows[0].cells[0].text

    def test_page_break_rendering(self):
        doc = self._make_doc_from_md("第一段\n\n<!-- pagebreak -->\n\n第二段")
        # Check for w:br type="page" in XML
        xml = str(doc.element.xml)
        assert "w:br" in xml and "page" in xml

    def test_bold_italic_runs(self):
        doc = self._make_doc_from_md("**加粗** 和 *斜体* 文本")
        normal_paras = [p for p in doc.paragraphs if p.style.name == "Normal"]
        assert len(normal_paras) >= 1
        # Find the specific run with "加粗" text and verify it's bold
        found_bold = False
        found_italic = False
        for p in normal_paras:
            for r in p.runs:
                if "加粗" in r.text and r.bold:
                    found_bold = True
                if "斜体" in r.text and r.italic:
                    found_italic = True
        assert found_bold, "Expected '加粗' text to be in a bold run"
        assert found_italic, "Expected '斜体' text to be in an italic run"

    def test_missing_image_fallback(self):
        doc = self._make_doc_from_md("![图表](./nonexistent.png)")
        # Should insert fallback text, not crash
        texts = [p.text for p in doc.paragraphs]
        assert any("图片加载失败" in t for t in texts)

    def test_full_document_structure(self):
        md = """# 测试报告

这是测试文档的正文。

## 数据概览

- 指标A: 100
- 指标B: 200

> 这是一条提示信息

## 数据表格

| 指标 | 数值 | 趋势 |
| --- | --- | --- |
| 收入 | 100万 | ↑ |
| 支出 | 50万 | ↓ |
"""
        doc = self._make_doc_from_md(md)
        assert len(doc.paragraphs) > 0
        assert len(doc.tables) >= 2  # 1 data table + 1 callout

    def test_save_to_file(self):
        doc = self._make_doc_from_md("# 测试\n\n保存测试")
        with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
            doc.save(f.name)
        assert os.path.exists(f.name)
        assert os.path.getsize(f.name) > 0
        os.unlink(f.name)

    def test_header_from_metadata(self):
        md = "# Test"
        spec = parse_markdown_to_docspec(md)
        spec["metadata"]["header"] = {"text": "页眉测试", "alignment": "center"}
        engine = DocxEngine()
        doc = engine.generate_from_spec(spec)
        section = doc.sections[0]
        header_text = section.header.paragraphs[0].text
        assert "页眉测试" in header_text

    def test_footer_page_number(self):
        md = "# Test"
        spec = parse_markdown_to_docspec(md)
        spec["metadata"]["footer"] = {"page_number": True}
        engine = DocxEngine()
        doc = engine.generate_from_spec(spec)
        section = doc.sections[0]
        footer_xml = section.footer._element.xml
        # Check for PAGE field code
        assert "instrText" in str(footer_xml) and "PAGE" in str(footer_xml)

    def test_table_header_styling_and_zebra(self):
        doc = self._make_doc_from_md("| 列1 | 列2 |\n| --- | --- |\n| A | B |\n| C | D |")
        assert len(doc.tables) == 1
        table = doc.tables[0]

        # Verify header cell has shading
        header_cell = table.rows[0].cells[0]
        tc_pr = header_cell._element.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd")
        assert len(tc_pr) > 0
        assert tc_pr[0].get("{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill") == "2563EB"

        # Verify zebra striping: row 2 (index 1, even data row) should NOT have shading,
        # row 3 (index 2, odd data row) SHOULD have shading
        row2_cell = table.rows[2].cells[0]  # second data row (r=1, odd -> zebra)
        tc_pr2 = row2_cell._element.findall(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}shd")
        assert len(tc_pr2) > 0

    def test_hr_bottom_border(self):
        doc = self._make_doc_from_md("第一段\n\n---\n\n第二段")
        # Find the HR paragraph (empty paragraph with pBdr/bottom)
        NS = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        for p in doc.paragraphs:
            pBdr = p._element.find("w:pPr/w:pBdr", NS)
            if pBdr is not None:
                bottom = pBdr.find("w:bottom", NS)
                if bottom is not None:
                    assert bottom.get(qn("w:val")) == "single"
                    assert bottom.get(qn("w:color")) == "CCCCCC"
                    return
        assert False, "No HR paragraph with bottom border found"

    def test_code_block_rendering(self):
        """代码块应渲染为表格（等宽字体+灰色背景）"""
        doc = self._make_doc_from_md("```python\nprint('hello')\n```")
        # Code block is rendered as a table
        assert len(doc.tables) == 1
        cell_text = doc.tables[0].rows[0].cells[0].text
        assert "python" in cell_text
        assert "print" in cell_text

    def test_code_block_without_language(self):
        """无语言标签的代码块也应正常渲染"""
        doc = self._make_doc_from_md("```\nsome code here\n```")
        assert len(doc.tables) == 1
        cell_text = doc.tables[0].rows[0].cells[0].text
        assert "some code here" in cell_text

    def test_toc_insertion(self):
        """metadata 含 toc 时应生成 TOC 域代码"""
        md = "# 标题一\n\n## 子标题"
        spec = parse_markdown_to_docspec(md)
        spec["metadata"]["toc"] = {"max_level": 3}
        engine = DocxEngine()
        doc = engine.generate_from_spec(spec)
        # TOC 插入应包含 "目录" 标题和 field code
        texts = [p.text for p in doc.paragraphs]
        assert "目录" in texts
        # 检查 field code
        xml = str(doc.element.xml)
        assert "TOC" in xml and "fldChar" in xml

    def test_toc_default_max_level(self):
        """toc 为 True 时默认 max_level 应为 3"""
        md = "# 标题一"
        spec = parse_markdown_to_docspec(md)
        spec["metadata"]["toc"] = True
        engine = DocxEngine()
        doc = engine.generate_from_spec(spec)
        texts = [p.text for p in doc.paragraphs]
        assert "目录" in texts

    def test_image_path_traversal_rejected(self):
        """路径穿越攻击应被拒绝"""
        doc = self._make_doc_from_md("![test](../../../etc/passwd)")
        texts = [p.text for p in doc.paragraphs]
        assert any("图片加载失败" in t for t in texts)

    def test_image_absolute_path_rejected(self):
        """绝对路径应被拒绝"""
        doc = self._make_doc_from_md("![test](/etc/passwd)")
        texts = [p.text for p in doc.paragraphs]
        assert any("图片加载失败" in t for t in texts)

    def test_image_valid_path_embedded(self):
        """有效图片路径应成功嵌入"""
        import base64
        # 创建一个 1x1 像素的 PNG 图片
        png_data = base64.b64decode(
            "iVBORw0KGgoAAAANSUhEUgAAAAEAAAABCAYAAAAfFcSJAAAACklEQVR4nGMAAQAABQABDQottAAAAABJRU5ErkJggg=="
        )
        # 在 CWD（项目根目录）创建，使用相对路径
        img_name = "test_pixel.png"
        img_path = os.path.join(os.getcwd(), img_name)
        with open(img_path, "wb") as f:
            f.write(png_data)

        try:
            md = f"![测试图片]({img_name})"
            spec = parse_markdown_to_docspec(md)
            engine = DocxEngine()
            doc = engine.generate_from_spec(spec)
            # 图片应嵌入，不应有 "图片加载失败"
            texts = [p.text for p in doc.paragraphs]
            assert not any("图片加载失败" in t for t in texts)
        finally:
            if os.path.exists(img_path):
                os.unlink(img_path)
